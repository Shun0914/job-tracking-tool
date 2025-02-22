import os
import shutil
from docx import Document

# スクリプトのディレクトリ
script_dir = os.path.dirname(os.path.abspath(__file__))
base_path = os.path.join(script_dir, "company")  # 企業フォルダ
master_path = os.path.join(script_dir, "master")  # マスターフォルダ

# マスターの職務経歴書と履歴書のパス
master_docx = os.path.join(master_path, "職務経歴書_master.docx")
master_pages = os.path.join(master_path, "履歴書_master.pages")

# 企業フォルダ一覧を取得
companies = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]

for company in companies:
    company_path = os.path.join(base_path, company)

    # 企業ごとの職務経歴書と履歴書のパス
    docx_path = os.path.join(company_path, f"職務経歴書_{company}.docx")
    pages_path = os.path.join(company_path, f"履歴書_{company}.pages")

    try:
        # 職務経歴書をコピー（上書きも許可）
        shutil.copy2(master_docx, docx_path)
        shutil.copy2(master_pages, pages_path)

        # コピー後のファイルサイズを確認
        docx_size = os.path.getsize(docx_path)
        pages_size = os.path.getsize(pages_path)

        print(f"✅ {company} の職務経歴書をコピー - サイズ: {docx_size} バイト")
        print(f"✅ {company} の履歴書をコピー - サイズ: {pages_size} バイト")

        # ファイルがコピーされているが、サイズが0Bの場合のエラーハンドリング
        if docx_size == 0 or pages_size == 0:
            print(f"⚠️ {company} のファイルサイズが0Bです。コピーが正常に行われていない可能性あり。")

        # コピー後の職務経歴書の内容をチェック
        doc = Document(docx_path)
        text = []

        # 通常の段落テキスト
        for para in doc.paragraphs:
            text.append(para.text)

        # テーブル内のテキストも取得
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)

        print(f"📄 {company} の職務経歴書の内容（最初の500文字）:")
        print("\n".join(text[:50]))  # 最初の50行を表示（長すぎるのを防ぐ）

    except FileNotFoundError:
        print(f"❌ {company} の職務経歴書のコピーに失敗: マスターファイルが存在しません ({master_docx})")
    except Exception as e:
        print(f"❌ {company} のコピー後の職務経歴書が開けません: {e}")

print("🎯 企業フォルダの書類をマスターから更新完了！")